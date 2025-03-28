import streamlit as st
import docx
import os
from fpdf import FPDF
import PyPDF2
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
from bs4 import BeautifulSoup
import requests
import yaml
import bcrypt
from pdf2image import convert_from_path
import io
from PIL import Image
import nltk
import textstat
from collections import Counter
import auth0_config


TEMPLATE_PATH = "templates/"
nltk.download("punkt")
st.markdown(
    """
    <style>
        /* Apply background color */
        body {
            background-color: #f4f7f9;
            font-family: 'Arial', sans-serif;
        }

        /* Sidebar styling */
        .sidebar .sidebar-content {
            background-color: #2C3E50;
            color: white;
        }

        /* Headings */
        h1, h2, h3 {
            color: #2C3E50;
            text-align: center;
            font-weight: bold;
        }

        /* Buttons */
        .stButton>button {
            background-color: #27ae60;
            color: white;
            border-radius: 8px;
            font-size: 16px;
            width: 100%;
            transition: 0.3s;
        }

        .stButton>button:hover {
            background-color: #2ecc71;
        }

        /* Cards for Resume Sections */
        .resume-section {
            background: white;
            border-radius: 10px;
            padding: 20px;
            margin-bottom: 20px;
            box-shadow: 0px 4px 10px rgba(0, 0, 0, 0.1);
        }

        /* Download Buttons */
        .stDownloadButton>button {
            background-color: #2980b9;
            color: white;
            border-radius: 8px;
            width: 100%;
            transition: 0.3s;
        }

        .stDownloadButton>button:hover {
            background-color: #3498db;
        }
    </style>
    """,
    unsafe_allow_html=True
)
if "theme_mode" not in st.session_state:
    st.session_state.theme_mode = "light"  # Default to light mode

# ✅ Toggle function
def toggle_theme():
    st.session_state.theme_mode = "dark" if st.session_state.theme_mode == "light" else "light"

# ✅ Sidebar: Theme Toggle Button
st.sidebar.markdown("### 🎨 Theme Settings")
st.sidebar.button("🌗 Toggle Dark/Light Mode", on_click=toggle_theme)

# ✅ Apply custom CSS based on theme selection
if st.session_state.theme_mode == "dark":
    st.markdown(
        """
        <style>
            body { background-color: #2C3E50; color: white; }
            .sidebar .sidebar-content { background-color: #1E2A38; color: white; }
            h1, h2, h3 { color: #F1C40F; }
            .stButton>button { background-color: #E74C3C; color: white; }
            .stButton>button:hover { background-color: #C0392B; }
        </style>
        """,
        unsafe_allow_html=True
    )
else:
    st.markdown(
        """
        <style>
            body { background-color: #F4F7F9; color: black; }
            .sidebar .sidebar-content { background-color: #FFFFFF; color: black; }
            h1, h2, h3 { color: #2C3E50; }
            .stButton>button { background-color: #27AE60; color: white; }
            .stButton>button:hover { background-color: #2ECC71; }
        </style>
        """,
        unsafe_allow_html=True
    )
    
def load_credentials():
    credentials_path = "credentials.yaml"

    # ✅ Check if the file exists
    if not os.path.exists(credentials_path):
        return {"credentials": {}}  # Return an empty dictionary if the file is missing

    with open(credentials_path, "r") as file:
        credentials = yaml.safe_load(file) or {}  # Ensure it never returns None
    
    # ✅ Ensure it has the expected structure
    if "credentials" not in credentials:
        credentials["credentials"] = {}

    return credentials

def save_credentials(credentials):
    with open("credentials.yaml", "w") as file:
        yaml.dump(credentials, file, default_flow_style=False)

# ✅ Hash passwords before saving
def hash_password(password):
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

# ✅ Function to check login
def authenticate(username, password):
    credentials = load_credentials()
    if username in credentials["credentials"]:
        stored_password = credentials["credentials"][username]["password"]
        return bcrypt.checkpw(password.encode(), stored_password.encode())
    return False

# ✅ Ensure session state variables are always initialized
if "extracted_skills" not in st.session_state:
    st.session_state.extracted_skills = []
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""

# ============================= LOGIN & SIGN-UP =============================
# Auth0 URLs
AUTH_URL = f"https://{auth0_config.AUTH0_DOMAIN}/authorize"
TOKEN_URL = f"https://{auth0_config.AUTH0_DOMAIN}/oauth/token"
login_url = f"{AUTH_URL}?response_type=code&client_id={auth0_config.CLIENT_ID}&redirect_uri={auth0_config.REDIRECT_URI}&scope=openid profile email"

if not st.session_state.logged_in:  # Show login options only if the user is not logged in
    st.title("🔐 Login or Sign Up")

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("Login"):
            st.session_state.auth_option = "login"

    with col2:
        if st.button("Sign Up"):
            st.session_state.auth_option = "signup"

    with col3:
        if st.button("Login with Google (Auth0)"):
            st.session_state.auth_option = "auth0"

    # Normal login form
    if st.session_state.get("auth_option") == "login":
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")

        if st.button("Submit Login"):
            if authenticate(username, password):
                st.session_state.logged_in = True
                st.session_state.username = username
                st.success("✅ Login successful!")
                st.experimental_rerun()
            else:
                st.error("❌ Invalid username or password!")

    # Sign-up form
    elif st.session_state.get("auth_option") == "signup":
        new_username = st.text_input("Choose a Username")
        new_password = st.text_input("Choose a Password", type="password")
        confirm_password = st.text_input("Confirm Password", type="password")

        if st.button("Create Account"):
            if new_password != confirm_password:
                st.error("❌ Passwords do not match!")
            else:
                credentials = load_credentials()
                if new_username in credentials["credentials"]:
                    st.error("❌ Username already exists! Choose another.")
                else:
                    credentials["credentials"][new_username] = {"password": hash_password(new_password)}
                    save_credentials(credentials)
                    st.success("✅ Account created successfully! Please log in.")

    # Auth0 Google Login
    elif st.session_state.get("auth_option") == "auth0":
        AUTH_URL = f"https://{auth0_config.AUTH0_DOMAIN}/authorize"
        login_url = f"{AUTH_URL}?response_type=code&client_id={auth0_config.CLIENT_ID}&redirect_uri={auth0_config.REDIRECT_URI}&scope=openid profile email"

        st.markdown(f"🔗 Click [here]({login_url}) to login with Google")

    st.stop()

# ✅ Logout button
st.sidebar.success(f"👋 Welcome, {st.session_state.username}")
if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.rerun()

# ✅ Set the correct Poppler path
POPPLER_PATH = r"C:\Users\HP\Documents\poppler-24.08.0\Library\bin"  # Update if needed

def convert_docx_to_image(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    image_path = docx_path.replace(".docx", ".png")

    # ✅ Convert DOCX to PDF
    os.system(f"libreoffice --headless --convert-to pdf {docx_path}")

    # ✅ Convert PDF to Image with Poppler path
    images = convert_from_path(pdf_path, poppler_path=POPPLER_PATH)

    if images:
        images[0].save(image_path, "PNG")  # Save first page as an image
        return image_path
    return None


# Function to replace placeholders in a template
def fill_template(template_path, user_data):
    doc = docx.Document(template_path)
    for para in doc.paragraphs:
        for key, value in user_data.items():
            if f"{{{key}}}" in para.text:
                para.text = para.text.replace(f"{{{key}}}", value)
    return doc
# Function to scrape jobs from Indeed
def scrape_indeed_jobs(query, num_jobs=5):
    base_url = f"https://www.indeed.com/jobs?q={query}"
    response = requests.get(base_url)
    soup = BeautifulSoup(response.text, "html.parser")
    job_list = []

    for job_card in soup.find_all("div", class_="job_seen_beacon")[:num_jobs]:
        title = job_card.find("h2").text.strip()
        company = job_card.find("span", class_="companyName").text.strip()
        job_link = "https://www.indeed.com" + job_card.find("a")["href"]
        job_list.append({"title": title, "company": company, "link": job_link})

    return job_list

# Function to scrape jobs from LinkedIn
def scrape_linkedin_jobs(query, num_jobs=5):
    base_url = f"https://www.linkedin.com/jobs/search?keywords={query}"
    response = requests.get(base_url)
    soup = BeautifulSoup(response.text, "html.parser")
    job_list = []

    for job_card in soup.find_all("div", class_="base-search-card")[:num_jobs]:
        title = job_card.find("h3").text.strip()
        company = job_card.find("h4").text.strip()
        job_link = job_card.find("a")["href"]
        job_list.append({"title": title, "company": company, "link": job_link})

    return job_list

# Function to scrape jobs from Glassdoor
def scrape_glassdoor_jobs(query, num_jobs=5):
    base_url = f"https://www.glassdoor.com/Job/jobs.htm?sc.keyword={query}"
    response = requests.get(base_url)
    soup = BeautifulSoup(response.text, "html.parser")
    job_list = []

    for job_card in soup.find_all("li", class_="react-job-listing")[:num_jobs]:
        title = job_card.find("a", class_="jobLink").text.strip()
        company = job_card.find("div", class_="jobHeader").text.strip()
        job_link = "https://www.glassdoor.com" + job_card.find("a")["href"]
        job_list.append({"title": title, "company": company, "link": job_link})

    return job_list

# Function to fetch jobs from multiple sources
def fetch_jobs_from_multiple_sources(query, num_jobs=5):
    jobs = []
    jobs.extend(scrape_indeed_jobs(query, num_jobs))
    jobs.extend(scrape_linkedin_jobs(query, num_jobs))
    jobs.extend(scrape_glassdoor_jobs(query, num_jobs))
    return jobs[:num_jobs]

# Function to match resume skills with job descriptions
def match_resumes_to_jobs(resume_skills, job_list):
    vectorizer = TfidfVectorizer()
    job_titles = [job["title"] for job in job_list]
    combined_texts = resume_skills + job_titles
    tfidf_matrix = vectorizer.fit_transform(combined_texts)
    resume_vector = tfidf_matrix[:len(resume_skills)]
    job_vectors = tfidf_matrix[len(resume_skills):]
    similarity_scores = cosine_similarity(resume_vector, job_vectors)
    return similarity_scores

# Function to extract text from PDF
def extract_text_from_pdf(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    return "\n".join([page.extract_text() for page in pdf_reader.pages])

def extract_skills(text):
    skills_list = [
        # Software Development
        "Python", "Java", "C++", "C#", "JavaScript", "TypeScript", "Go", "Swift", "Kotlin", "Ruby", "PHP", "Rust",
        "HTML", "CSS", "React", "Angular", "Vue.js", "Node.js", "Django", "Flask", "Spring Boot", "Express.js",
        "Git", "Docker", "Kubernetes", "CI/CD", "Jenkins", "GraphQL", "REST API", "SOAP", "Microservices",
        
        # Data Science & AI
        "Machine Learning", "Deep Learning", "Data Science", "Artificial Intelligence", "NLP", "Computer Vision",
        "TensorFlow", "PyTorch", "Scikit-Learn", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Keras", "OpenCV",
        "Big Data", "Hadoop", "Spark", "Apache Kafka", "ETL", "Tableau", "Power BI", "Data Visualization",
        
        # Cloud Computing
        "AWS", "Azure", "Google Cloud", "DevOps", "Cloud Security", "Terraform", "Ansible", "CloudFormation",
        "Serverless", "Lambda", "EC2", "S3", "Kubernetes", "Cloud Networking",
        
        # Cybersecurity
        "Ethical Hacking", "Penetration Testing", "Malware Analysis", "Network Security", "SIEM", "SOC", 
        "Firewall Management", "Incident Response", "Cryptography", "Zero Trust Security", "Identity Management",
        
        # Business & Management
        "Agile", "Scrum", "Kanban", "Project Management", "Business Analysis", "Risk Management",
        "Stakeholder Management", "Product Management", "Lean Methodology", "Six Sigma",
        
        # Digital Marketing
        "SEO", "SEM", "Google Ads", "Facebook Ads", "Social Media Marketing", "Email Marketing",
        "Marketing Automation", "Content Strategy", "Copywriting", "PPC", "Google Analytics",
        
        # Finance & Accounting
        "Financial Analysis", "Accounting", "Budgeting", "Forecasting", "Investment Analysis",
        "Risk Assessment", "Auditing", "Taxation", "Excel", "QuickBooks", "SAP",
        
        # Networking & IT Support
        "Network Security", "CCNA", "CCNP", "Routing", "Switching", "LAN", "WAN", "VPN", "TCP/IP",
        "Linux Administration", "Windows Server", "Active Directory",
        
        # Soft Skills
        "Communication", "Leadership", "Time Management", "Problem Solving", "Critical Thinking",
        "Teamwork", "Adaptability", "Creativity"
    ]
    
    return list(set([word for word in re.findall(r'\b\w+\b', text) if word.lower() in [skill.lower() for skill in skills_list]]))


# ✅ Function to calculate Resume Score
def score_resume(resume_text, extracted_skills):
    score = 0
    total_factors = 5  # Total factors for scoring

    # 1️⃣ **Keyword Matching (20%)**
    industry_keywords = ["data analysis", "machine learning", "business intelligence", "python", "sql", "power bi", "dashboard", "visualization"]
    matched_keywords = sum(1 for word in industry_keywords if word.lower() in resume_text.lower())
    keyword_score = (matched_keywords / len(industry_keywords)) * 20

    # 2️⃣ **Experience Level (20%)**
    experience_data = extract_experience(resume_text)  # ✅ Get experience text & years
    experience_years = experience_data["years_of_experience"]  # ✅ Extract only years
    experience_score = 20 if experience_years >= 3 else 10 if experience_years >= 1 else 5  # ✅ Fix TypeError

    # 3️⃣ **Readability Score (20%)**
    readability = textstat.flesch_reading_ease(resume_text)
    readability_score = 20 if readability > 50 else 10 if readability > 30 else 5

    # 4️⃣ **Skill Relevance (20%)**
    matched_skills = len(set(extracted_skills).intersection(set(industry_keywords)))
    skill_score = (matched_skills / len(industry_keywords)) * 20

    # 5️⃣ **ATS Compliance (20%)**
    ats_keywords = ["education", "experience", "skills", "projects", "certifications"]
    ats_score = sum(1 for word in ats_keywords if word.lower() in resume_text.lower())
    ats_score = (ats_score / len(ats_keywords)) * 20

    # ✅ Calculate Final Score (Out of 100)
    final_score = keyword_score + experience_score + readability_score + skill_score + ats_score
    return round(final_score, 2)


def display_resume_preview(user_data):
    st.markdown("### 📄 Live Resume Preview")
    st.markdown("---")
    st.markdown(f"**👤 Name:** {user_data.get('NAME', 'Your Name')}")
    st.markdown(f"📧 **Email:** {user_data.get('EMAIL', 'your.email@example.com')}")
    st.markdown(f"📞 **Phone:** {user_data.get('PHONE', '123-456-7890')}")
    st.markdown(f"🎓 **Education:** {user_data.get('EDUCATION', 'Your Degree, University Name')}")
    st.markdown(f"💼 **Experience:** {user_data.get('EXPERIENCE', 'Your work experience details...')}")
    st.markdown(f"🛠 **Skills:** {user_data.get('SKILLS', 'List your skills...')}")
    st.markdown(f"🏆 **Certifications:** {user_data.get('CERTIFICATIONS', 'Your certifications...')}")
    st.markdown("---")
    
import re

import re

def extract_experience(text):
    """
    Extracts the full experience section and calculates years of experience.

    Parameters:
    text (str): Resume text.

    Returns:
    dict: {
        "experience_text": Full extracted experience section or "Experience not found",
        "years_of_experience": Number of years found (int, default 0)
    }
    """

    # ✅ Define common headers that indicate work experience sections
    experience_headers = [
        "Work Experience/Internships", "work experience", "professional experience", 
        "employment history", "career summary", "job experience", "experience"
    ]

    # ✅ Regex pattern to detect the experience section
    experience_pattern = r"(?:{})\s*(?:[:\-]?\s*)\n?(.*?)(?:\n\s*\n|\Z)".format("|".join(experience_headers))
    matches = re.findall(experience_pattern, text, re.IGNORECASE | re.DOTALL)

    extracted_experience = matches[0].strip() if matches else "Experience not found"

    # ✅ Extract number of years from the experience section
    years_pattern = r"(\d+)\s*(?:years?|yrs?)"
    years_match = re.findall(years_pattern, extracted_experience, re.IGNORECASE)

    years_of_experience = max(map(int, years_match)) if years_match else 0  # Take the highest year found

    return {
        "experience_text": extracted_experience,
        "years_of_experience": years_of_experience
    }

# ✅ Example Usage:
resume_text = """ 
Work Experience
Software Engineer at XYZ Corp (2018 - Present)
- Developed machine learning models for fraud detection.
- Optimized SQL queries to improve database performance.
- Led a team of 5 developers for AI-driven automation projects.
Worked as Data Analyst for 3 years in ABC Ltd.
"""
extracted_data = extract_experience(resume_text)
print("Extracted Experience Section:\n", extracted_data["experience_text"])
print("Extracted Years of Experience:", extracted_data["years_of_experience"])

# Function to convert DOCX to PDF
def convert_docx_to_pdf(doc, pdf_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for para in doc.paragraphs:
        pdf.multi_cell(200, 10, txt=para.text.encode("latin-1", "ignore").decode("latin-1"), align='L')
    pdf.output(pdf_path)

if not os.path.exists(TEMPLATE_PATH):
    os.makedirs(TEMPLATE_PATH)

# ✅ Sidebar Navigation
st.sidebar.title("🔍 Navigation")
option = st.sidebar.radio("Choose an option:", ["Resume Analyzer", "Get Job Recommendations", "Resume Generator"])
if option == "Resume Generator":
    st.title("📝 Generate your own resume")

elif option == "Resume Analyzer":
    st.title("🛠️ Skills Extractor")

elif option == "Get Job Recommendations":
    st.title("🔍 Get Job Recommendations")

if option == "Resume Generator":
    st.subheader("📝 Create a Resume Using Templates")
    templates = [f for f in os.listdir(TEMPLATE_PATH) if f.endswith(".docx")]
    if templates:
        selected_template = st.selectbox("Select a Resume Template", templates)
        
        # ✅ Show Template Preview
        template_path = os.path.join(TEMPLATE_PATH, selected_template)
        # preview_image = convert_docx_to_image(template_path)
        
        # if preview_image:
        #    st.image(preview_image, caption="Template Preview", use_column_width=True)
        # else:
          #  st.warning("⚠️ Could not generate a preview for this template.")

        user_data = {key: st.text_input(key) for key in ["NAME", "EMAIL", "PHONE", "SKILLS", "EXPERIENCE", "EDUCATION", "CERTIFICATIONS"]}
        
        # Display Live Resume Preview
        display_resume_preview(user_data)

        if st.button("Generate Resume"):
            updated_doc = fill_template(template_path, user_data)
            docx_path = "Generated_Resume.docx"
            updated_doc.save(docx_path)

            pdf_path = "Generated_Resume.pdf"
            convert_docx_to_pdf(updated_doc, pdf_path)

            # ✅ Fix file download by using BytesIO
            with open(docx_path, "rb") as docx_file:
                docx_bytes = io.BytesIO(docx_file.read())
                st.download_button("Download Resume (DOCX)", docx_bytes, file_name="Generated_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            with open(pdf_path, "rb") as pdf_file:
                pdf_bytes = io.BytesIO(pdf_file.read())
                st.download_button("Download Resume (PDF)", pdf_bytes, file_name="Generated_Resume.pdf", mime="application/pdf")
    else:
        st.warning("No resume templates found! Upload DOCX templates in 'templates/' folder.")


elif option == "Resume Analyzer":
    uploaded_file = st.file_uploader("Upload Your Resume (PDF)", type=["pdf"])
    
    if uploaded_file:
        resume_text = extract_text_from_pdf(uploaded_file)
        extracted_skills = extract_skills(resume_text)
        extracted_experience = extract_experience(resume_text)

        st.write("**Extracted Skills:**", extracted_skills)
        st.write("**Extracted Experience:**", extracted_experience)

        # ✅ Calculate Resume Score
        resume_score = score_resume(resume_text, extracted_skills)
        st.subheader(f"📊 Your Resume Score: {resume_score}/100")

        # ✅ Show Score Interpretation
        if resume_score >= 80:
            st.success("🔥 Excellent Resume! Ready for job applications.")
        elif resume_score >= 60:
            st.warning("⚠️ Good Resume! Consider optimizing your skills section.")
        else:
            st.error("❌ Your Resume needs improvement. Add more industry keywords & details.")


# ============================= JOB RECOMMENDATIONS =============================
elif option == "Get Job Recommendations":
    uploaded_file = st.file_uploader("Upload Your Resume (PDF) to Get Job Recommendations", type=["pdf"])

    if uploaded_file:
        resume_text = extract_text_from_pdf(uploaded_file)
        extracted_skills = extract_skills(resume_text)

        # ✅ Store extracted skills in session state
        st.session_state.extracted_skills = extracted_skills

        st.write("**Extracted Skills:**", extracted_skills)

    # ✅ Ensure extracted skills exist before proceeding
    if st.session_state.extracted_skills:
        jobs = fetch_jobs_from_multiple_sources(",".join(st.session_state.extracted_skills), num_jobs=10)

        st.subheader("🔍 Recommended Job Listings from Multiple Websites")
        if jobs:
            for job in jobs:
                st.write(f"**{job['title']}** at **{job['company']}**")
                st.markdown(f"[Apply Here]({job['link']})", unsafe_allow_html=True)
        else:
            st.warning("No jobs found based on your skills.")
    else:
        st.warning("⚠️ Please upload a resume to get job recommendations.")


